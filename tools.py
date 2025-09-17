from RAG import RAGPipeline
import requests
from bs4 import BeautifulSoup


# Store global RAG instance
rag_pipeline = RAGPipeline()

def get_website_content(url: str) -> str:
    
    resp = requests.get(url, timeout=20)
    resp.raise_for_status()
    soup = BeautifulSoup(resp.text, "html.parser")

    for s in soup(["script", "style", "noscript"]):
        s.extract()

    text = soup.get_text(separator=" ", strip=True)

    # Build RAG index from extracted text
    rag_pipeline.build_index(text, source=url)

    return text

def retrieve_relevant_context(query: str, k: int = 5) -> str:
    return rag_pipeline.get_context(query, k=k)


def generate_word_report(report_summary):
    from docx import Document
    import os

    doc = Document()
    doc.add_heading(report_summary.title, 0)
    doc.add_paragraph(f"Date: {report_summary.date}")
    doc.add_paragraph(report_summary.description)

    # Prices
    doc.add_heading("Prices", level=1)
    for p in report_summary.prices:
        doc.add_paragraph(f"- {p.energy}: {p.price} ({p.evolution})")

    # Save file
    output_path = os.path.abspath("reports/report.docx")
    doc.save(output_path)
    return output_path
